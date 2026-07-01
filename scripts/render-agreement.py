"""Render an agreement document by replacing Dhivehi placeholders.

Uses python-docx to merge split XML runs before doing find-and-replace,
which solves the issue of Word fragmenting placeholder text across runs.

Usage:
    echo '{"placeholder": "value"}' | python3 render-agreement.py \
        --template path/to/template.docx \
        --output-dir /tmp/agreegen-xyz
"""

import sys
import json
import argparse
from pathlib import Path

from docx import Document


def merge_runs(paragraph):
    """Merge adjacent runs that share identical formatting into one run.

    Word/LibreOffice often splits contiguous text into multiple <w:r> elements,
    especially with RTL scripts. This merges them back so find-and-replace
    can match placeholders that span multiple runs.
    """
    runs = paragraph.runs
    if len(runs) <= 1:
        return

    merged = [runs[0]]
    for run in runs[1:]:
        prev = merged[-1]
        if _same_formatting(prev, run):
            prev.text += run.text
            # Clear the consumed run so it doesn't render
            run._element.getparent().remove(run._element)
        else:
            merged.append(run)


def _same_formatting(a, b):
    """Check if two runs share the same formatting properties."""
    a_props = a._element.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
    )
    b_props = b._element.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
    )

    if a_props is None and b_props is None:
        return True
    if a_props is None or b_props is None:
        return False

    from lxml import etree

    return etree.tostring(a_props) == etree.tostring(b_props)


def replace_in_paragraphs(paragraphs, replacements):
    """Merge runs then apply replacements across all paragraphs."""
    for paragraph in paragraphs:
        merge_runs(paragraph)
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)


def replace_in_doc(doc, replacements):
    """Apply replacements in body paragraphs, tables, headers, and footers."""
    # Body paragraphs
    replace_in_paragraphs(doc.paragraphs, replacements)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, replacements)

    # Headers and footers
    for section in doc.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            if header and header.is_linked_to_previous is False:
                replace_in_paragraphs(header.paragraphs, replacements)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_paragraphs(cell.paragraphs, replacements)

        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if footer and footer.is_linked_to_previous is False:
                replace_in_paragraphs(footer.paragraphs, replacements)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_paragraphs(cell.paragraphs, replacements)


def main():
    parser = argparse.ArgumentParser(
        description="Render agreement document from template"
    )
    parser.add_argument("--template", required=True, help="Path to the .docx template")
    parser.add_argument(
        "--output-dir", required=True, help="Directory to write output.docx"
    )
    args = parser.parse_args()

    template_path = Path(args.template)
    if not template_path.exists():
        print(f"template not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    output_dir = Path(args.output_dir)
    if not output_dir.is_dir():
        print(f"output directory does not exist: {output_dir}", file=sys.stderr)
        sys.exit(1)

    try:
        replacements = json.load(sys.stdin)
    except json.JSONDecodeError as e:
        print(f"invalid JSON on stdin: {e}", file=sys.stderr)
        sys.exit(1)

    if not isinstance(replacements, dict):
        print(
            "stdin JSON must be an object mapping placeholders to values",
            file=sys.stderr,
        )
        sys.exit(1)

    doc = Document(str(template_path))
    replace_in_doc(doc, replacements)

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))


if __name__ == "__main__":
    main()
